"""SharePoint ingestion via Microsoft Graph API (app-only / client credentials).

Each configured SharePoint URL is treated as a project root. We walk every
folder under it recursively, download each supported file, extract its text,
and emit a Document tagged with:
    - project: derived from the URL's last path segment
    - source:  the SharePoint webUrl (clickable link)
    - path:    the file's path within the drive
    - type:    "sharepoint_<ext>"
"""

from __future__ import annotations

import io
import logging
from dataclasses import dataclass
from typing import Iterable
from urllib.parse import unquote, urlparse

import requests
from bs4 import BeautifulSoup
from langchain_core.documents import Document
from pypdf import PdfReader

logger = logging.getLogger(__name__)

GRAPH_BASE = "https://graph.microsoft.com/v1.0"
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".txt", ".md", ".markdown", ".html", ".htm", ".csv"}
# SharePoint document-library display names that map to a Graph drive. We try
# the URL segment as-is first; if it doesn't match, we fall back to the site's
# default drive.
KNOWN_DRIVE_ALIASES = {
    "shared documents": "Documents",
    "documents": "Documents",
}


class SharePointLoadError(RuntimeError):
    pass


@dataclass
class _Target:
    """Parsed parts of a SharePoint URL pointing at a folder."""
    hostname: str
    site_relative_path: str  # e.g. "/sites/MySite"
    drive_segment: str | None  # e.g. "Shared Documents" or None
    folder_path: str  # e.g. "/MyProjectFolder/Sub" (may be "")
    project_name: str


def load_sharepoint_documents(
    urls: list[str],
    tenant_id: str,
    client_id: str,
    client_secret: str,
) -> list[Document]:
    if not urls:
        return []
    if not (tenant_id and client_id and client_secret):
        raise SharePointLoadError(
            "SharePoint URLs configured but SHAREPOINT_TENANT_ID / "
            "SHAREPOINT_CLIENT_ID / SHAREPOINT_CLIENT_SECRET are missing."
        )

    token = _acquire_token(tenant_id, client_id, client_secret)
    session = requests.Session()
    session.headers.update(
        {
            "Authorization": f"Bearer {token}",
            "Accept": "application/json",
            "User-Agent": "drs-chatbot-sharepoint/1.0",
        }
    )

    documents: list[Document] = []
    for raw_url in urls:
        try:
            documents.extend(_load_one_root(session, raw_url))
        except SharePointLoadError as exc:
            logger.warning("Skipping SharePoint root '%s': %s", raw_url, exc)
        except Exception as exc:  # pragma: no cover - defensive
            logger.warning("Unexpected error on SharePoint root '%s': %s", raw_url, exc)
    return documents


def _acquire_token(tenant_id: str, client_id: str, client_secret: str) -> str:
    # MSAL is the supported path; fall back to a direct token call if the
    # library happens to be missing so the rest of the bot still imports cleanly.
    try:
        from msal import ConfidentialClientApplication
    except ImportError as exc:  # pragma: no cover
        raise SharePointLoadError(
            "msal is not installed. Run `pip install msal` (it's in requirements.txt)."
        ) from exc

    app = ConfidentialClientApplication(
        client_id=client_id,
        client_credential=client_secret,
        authority=f"https://login.microsoftonline.com/{tenant_id}",
    )
    result = app.acquire_token_for_client(scopes=["https://graph.microsoft.com/.default"])
    if "access_token" not in result:
        raise SharePointLoadError(
            "Failed to acquire Microsoft Graph token. "
            f"{result.get('error')}: {result.get('error_description')}"
        )
    return result["access_token"]


def _load_one_root(session: requests.Session, raw_url: str) -> list[Document]:
    target = _parse_url(raw_url)
    site_id = _resolve_site_id(session, target.hostname, target.site_relative_path)
    drive_id, root_item_id = _resolve_folder(session, site_id, target)

    documents: list[Document] = []
    _walk(
        session=session,
        drive_id=drive_id,
        item_id=root_item_id,
        project_name=target.project_name,
        documents=documents,
    )
    return documents


def _parse_url(raw_url: str) -> _Target:
    parsed = urlparse(raw_url.strip())
    if not parsed.scheme or not parsed.netloc:
        raise SharePointLoadError(f"Invalid SharePoint URL: {raw_url!r}")

    hostname = parsed.netloc
    path_segments = [unquote(seg) for seg in parsed.path.split("/") if seg]
    if not path_segments:
        raise SharePointLoadError(f"URL has no path: {raw_url!r}")

    # Detect a site collection prefix (sites/<name> or teams/<name>). Without
    # one we treat the path as living on the root site.
    if path_segments[0].lower() in {"sites", "teams"} and len(path_segments) >= 2:
        site_relative_path = f"/{path_segments[0]}/{path_segments[1]}"
        remainder = path_segments[2:]
    else:
        site_relative_path = "/"
        remainder = path_segments

    drive_segment = remainder[0] if remainder else None
    folder_segments = remainder[1:] if remainder else []
    folder_path = "/" + "/".join(folder_segments) if folder_segments else ""
    project_name = (folder_segments[-1] if folder_segments else (drive_segment or path_segments[-1])).strip()

    return _Target(
        hostname=hostname,
        site_relative_path=site_relative_path,
        drive_segment=drive_segment,
        folder_path=folder_path,
        project_name=project_name,
    )


def _resolve_site_id(session: requests.Session, hostname: str, site_path: str) -> str:
    suffix = "" if site_path == "/" else f":{site_path}"
    url = f"{GRAPH_BASE}/sites/{hostname}{suffix}"
    response = session.get(url, timeout=30)
    if response.status_code == 404:
        raise SharePointLoadError(
            f"SharePoint site not found: {hostname}{site_path}. "
            "Check the URL and that the AAD app has Sites.Read.All granted."
        )
    _raise_for_graph(response, context=f"resolving site {hostname}{site_path}")
    payload = response.json()
    site_id = payload.get("id")
    if not site_id:
        raise SharePointLoadError(f"Site response missing id: {payload}")
    return site_id


def _resolve_folder(
    session: requests.Session,
    site_id: str,
    target: _Target,
) -> tuple[str, str]:
    """Resolve to (drive_id, folder_item_id) for the configured URL."""

    candidate_drives = _list_candidate_drives(session, site_id, target.drive_segment)
    last_error: SharePointLoadError | None = None
    for drive in candidate_drives:
        drive_id = drive["id"]
        try:
            item_id = _resolve_folder_in_drive(session, drive_id, target.folder_path)
        except SharePointLoadError as exc:
            last_error = exc
            continue
        return drive_id, item_id
    if last_error:
        raise last_error
    raise SharePointLoadError(
        f"Could not resolve folder for {target.site_relative_path}{target.drive_segment or ''}"
        f"{target.folder_path}"
    )


def _list_candidate_drives(
    session: requests.Session,
    site_id: str,
    drive_segment: str | None,
) -> list[dict]:
    response = session.get(f"{GRAPH_BASE}/sites/{site_id}/drives", timeout=30)
    _raise_for_graph(response, context=f"listing drives for site {site_id}")
    drives = response.json().get("value", []) or []
    if not drive_segment:
        return drives

    normalized = drive_segment.strip().lower()
    aliased = KNOWN_DRIVE_ALIASES.get(normalized, drive_segment).lower()

    matched: list[dict] = []
    other: list[dict] = []
    for drive in drives:
        name = (drive.get("name") or "").strip().lower()
        if name in {normalized, aliased}:
            matched.append(drive)
        else:
            other.append(drive)
    # Try matched drives first, then the default drive, then everything else.
    if not matched:
        return drives
    return matched + other


def _resolve_folder_in_drive(
    session: requests.Session,
    drive_id: str,
    folder_path: str,
) -> str:
    if not folder_path or folder_path == "/":
        url = f"{GRAPH_BASE}/drives/{drive_id}/root"
    else:
        # Graph wants `:/path:` form for path-based addressing.
        path = folder_path if folder_path.startswith("/") else f"/{folder_path}"
        url = f"{GRAPH_BASE}/drives/{drive_id}/root:{path}"
    response = session.get(url, timeout=30)
    if response.status_code == 404:
        raise SharePointLoadError(f"Folder not found in drive: {folder_path}")
    _raise_for_graph(response, context=f"resolving folder {folder_path} in drive {drive_id}")
    item = response.json()
    item_id = item.get("id")
    if not item_id:
        raise SharePointLoadError(f"Folder response missing id: {item}")
    return item_id


def _walk(
    session: requests.Session,
    drive_id: str,
    item_id: str,
    project_name: str,
    documents: list[Document],
) -> None:
    next_url = f"{GRAPH_BASE}/drives/{drive_id}/items/{item_id}/children?$top=200"
    while next_url:
        response = session.get(next_url, timeout=30)
        _raise_for_graph(response, context=f"listing children of item {item_id}")
        payload = response.json()
        for child in payload.get("value", []) or []:
            if "folder" in child:
                _walk(
                    session=session,
                    drive_id=drive_id,
                    item_id=child["id"],
                    project_name=project_name,
                    documents=documents,
                )
            elif "file" in child:
                doc = _ingest_file(session, child, project_name)
                if doc is not None:
                    documents.append(doc)
        next_url = payload.get("@odata.nextLink")


def _ingest_file(
    session: requests.Session,
    item: dict,
    project_name: str,
) -> Document | None:
    name = item.get("name") or ""
    ext = _extension(name)
    if ext not in SUPPORTED_EXTENSIONS:
        return None

    download_url = item.get("@microsoft.graph.downloadUrl")
    if not download_url:
        logger.warning("File missing downloadUrl, skipping: %s", name)
        return None

    try:
        # downloadUrl is a pre-authenticated short-lived URL — do NOT send
        # the Bearer token, that would fail with 400.
        content = requests.get(download_url, timeout=60).content
    except Exception as exc:  # pragma: no cover - defensive
        logger.warning("Failed to download SharePoint file %s: %s", name, exc)
        return None

    text = _extract_text(name, ext, content)
    if not text:
        return None

    return Document(
        page_content=text,
        metadata={
            "source": item.get("webUrl") or name,
            "path": _parent_path(item),
            "project": project_name,
            "type": f"sharepoint_{ext.lstrip('.')}",
        },
    )


def _extension(name: str) -> str:
    if "." not in name:
        return ""
    return "." + name.rsplit(".", 1)[1].lower()


def _parent_path(item: dict) -> str:
    parent = item.get("parentReference") or {}
    parent_path = parent.get("path") or ""
    # parent.path is like "/drives/{driveId}/root:/FolderA/SubFolder"
    if ":" in parent_path:
        parent_path = parent_path.split(":", 1)[1]
    return f"{parent_path}/{item.get('name', '')}".lstrip("/")


def _extract_text(name: str, ext: str, content: bytes) -> str:
    try:
        if ext == ".pdf":
            return _extract_pdf(content)
        if ext == ".docx":
            return _extract_docx(content)
        if ext == ".xlsx":
            return _extract_xlsx(content)
        if ext in {".html", ".htm"}:
            return _extract_html(content)
        # txt / md / csv — best-effort utf-8
        return _clean_text(content.decode("utf-8", errors="ignore"))
    except Exception as exc:  # pragma: no cover - defensive
        logger.warning("Failed to extract text from SharePoint file %s: %s", name, exc)
        return ""


def _extract_pdf(content: bytes) -> str:
    reader = PdfReader(io.BytesIO(content))
    chunks: list[str] = []
    for page in reader.pages:
        chunks.append(page.extract_text() or "")
    return _clean_text("\n".join(chunks))


def _extract_docx(content: bytes) -> str:
    from docx import Document as DocxDocument  # python-docx

    doc = DocxDocument(io.BytesIO(content))
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        if paragraph.text and paragraph.text.strip():
            parts.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return _clean_text("\n".join(parts))


def _extract_xlsx(content: bytes) -> str:
    from openpyxl import load_workbook

    workbook = load_workbook(io.BytesIO(content), data_only=True, read_only=True)
    parts: list[str] = []
    for sheet in workbook.worksheets:
        parts.append(f"# Sheet: {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            cells = [str(cell).strip() for cell in row if cell is not None and str(cell).strip()]
            if cells:
                parts.append(" | ".join(cells))
    return _clean_text("\n".join(parts))


def _extract_html(content: bytes) -> str:
    soup = BeautifulSoup(content, "html.parser")
    for tag in soup(["script", "style", "noscript"]):
        tag.extract()
    return _clean_text(soup.get_text("\n"))


def _clean_text(text: str) -> str:
    lines = [line.strip() for line in (text or "").splitlines()]
    return "\n".join(line for line in lines if line)


def _raise_for_graph(response: requests.Response, context: str) -> None:
    if response.status_code < 400:
        return
    body_snippet = (response.text or "")[:300]
    raise SharePointLoadError(
        f"Microsoft Graph error while {context}: HTTP {response.status_code}. {body_snippet}"
    )
