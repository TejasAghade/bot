from __future__ import annotations

import base64
import io
import logging
from pathlib import Path
import re
from urllib.parse import quote, urlparse

import requests
from bs4 import BeautifulSoup
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from pypdf import PdfReader

logger = logging.getLogger(__name__)

# Safeguards against pathological spreadsheets. openpyxl will happily iterate
# over hundreds of thousands of phantom rows (formatting/data left in a huge
# used-range), which stalls ingestion for hours. Cap the work per file.
MAX_XLSX_ROWS_PER_SHEET = 20_000
MAX_XLSX_CHARS = 2_000_000

SUPPORTED_EXTENSIONS = {
    ".txt",
    ".md",
    ".markdown",
    ".pdf",
    ".html",
    ".htm",
    ".docx",
    ".xlsx",
}


class DocumentLoadError(RuntimeError):
    pass


class IncrementalSink:
    """Persist documents one at a time so ingestion is interruptible and resumable.

    ``persist`` is called with each newly loaded Document (it chunks the document and
    writes it to the vector store immediately). ``ingested`` holds the ``source`` keys
    already present in the store, so an interrupted run can be resumed with
    ``append=True`` without reprocessing files that finished earlier.
    """

    def __init__(self, persist, ingested=None):
        self._persist = persist
        self.ingested: set[str] = set(ingested or ())

    def is_ingested(self, source: str) -> bool:
        return bool(source) and source in self.ingested

    def emit(self, doc: Document) -> None:
        source = str(doc.metadata.get("source", ""))
        if self.is_ingested(source):
            return
        self._persist(doc)
        if source:
            self.ingested.add(source)


def load_documents(
    data_dir: str,
    urls_file: str | None = None,
    azure_devops_pat: str | None = None,
    azure_devops_org: str | None = None,
    azure_devops_project: str | None = None,
    azure_devops_projects: list[str] | None = None,
    azure_devops_wiki: str | None = None,
    azure_devops_wiki_path: str = "/",
    azure_devops_api_version: str = "7.1",
    sharepoint_tenant_id: str | None = None,
    sharepoint_client_id: str | None = None,
    sharepoint_client_secret: str | None = None,
    sharepoint_site: str | None = None,
    sharepoint_folders: list[str] | None = None,
    sharepoint_graph_base_url: str = "https://graph.microsoft.com/v1.0",
    sharepoint_login_base_url: str = "https://login.microsoftonline.com",
    sink: IncrementalSink | None = None,
) -> list[Document]:
    # When a sink is provided, each document is persisted the moment it is loaded
    # (interruptible/resumable ingestion). Otherwise documents are collected into a
    # list and returned, preserving the original batch API for other callers/tests.
    collected: list[Document] = []

    def emit(docs: list[Document]) -> None:
        for doc in docs:
            if sink is not None:
                sink.emit(doc)
            else:
                collected.append(doc)

    emit(load_local_documents(data_dir))
    if urls_file:
        emit(load_url_documents(urls_file, azure_devops_pat=azure_devops_pat))

    projects = _resolve_project_list(azure_devops_projects, azure_devops_project)
    if azure_devops_pat and azure_devops_org and projects:
        # When multiple projects are configured, ignore AZURE_DEVOPS_WIKI (a single
        # wiki identifier cannot meaningfully apply across projects) and ingest every
        # wiki in each project.
        restrict_to_single_wiki = bool(azure_devops_wiki) and len(projects) == 1
        for project_name in projects:
            try:
                if restrict_to_single_wiki:
                    emit(
                        load_azure_devops_wiki_documents(
                            organization=azure_devops_org,
                            project=project_name,
                            wiki_identifier=azure_devops_wiki,
                            wiki_path=azure_devops_wiki_path,
                            azure_devops_pat=azure_devops_pat,
                            api_version=azure_devops_api_version,
                        )
                    )
                else:
                    emit(
                        load_all_azure_devops_project_wikis(
                            organization=azure_devops_org,
                            project=project_name,
                            wiki_path=azure_devops_wiki_path,
                            azure_devops_pat=azure_devops_pat,
                            api_version=azure_devops_api_version,
                        )
                    )
            except DocumentLoadError as exc:
                logger.warning("Skipping Azure DevOps project '%s': %s", project_name, exc)

    if sharepoint_tenant_id and sharepoint_client_id and sharepoint_client_secret:
        try:
            sharepoint_docs = load_sharepoint_documents(
                tenant_id=sharepoint_tenant_id,
                client_id=sharepoint_client_id,
                client_secret=sharepoint_client_secret,
                site=sharepoint_site,
                folders=sharepoint_folders,
                graph_base_url=sharepoint_graph_base_url,
                login_base_url=sharepoint_login_base_url,
                sink=sink,
            )
            # With a sink, SharePoint files are emitted per-file as they load (so a
            # pause keeps them); the return value is empty in that case.
            if sink is None:
                collected.extend(sharepoint_docs)
        except DocumentLoadError as exc:
            logger.warning("Skipping SharePoint ingestion: %s", exc)
    return collected


def _resolve_project_list(
    projects: list[str] | None,
    fallback_project: str | None,
) -> list[str]:
    resolved: list[str] = []
    seen: set[str] = set()
    for name in projects or []:
        cleaned = (name or "").strip()
        if cleaned and cleaned not in seen:
            seen.add(cleaned)
            resolved.append(cleaned)
    if not resolved and fallback_project:
        cleaned = fallback_project.strip()
        if cleaned:
            resolved.append(cleaned)
    return resolved


def load_local_documents(data_dir: str) -> list[Document]:
    base = Path(data_dir)
    if not base.exists():
        logger.warning("Data directory does not exist: %s", base)
        return []

    documents: list[Document] = []
    for path in base.rglob("*"):
        if not path.is_file():
            continue

        if path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            continue

        try:
            documents.extend(_load_single_file(path))
        except Exception as exc:  # pragma: no cover - defensive logging
            logger.warning("Failed to load file %s: %s", path, exc)
    return documents


def load_url_documents(urls_file: str, azure_devops_pat: str | None = None) -> list[Document]:
    path = Path(urls_file)
    if not path.exists():
        return []

    urls = [
        line.strip()
        for line in path.read_text(encoding="utf-8").splitlines()
        if line.strip() and not line.strip().startswith("#")
    ]
    documents: list[Document] = []
    session = requests.Session()
    session.headers.update(
        {
            "User-Agent": (
                "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                "AppleWebKit/537.36 (KHTML, like Gecko) "
                "Chrome/133.0.0.0 Safari/537.36"
            ),
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
            "Accept-Language": "en-US,en;q=0.8",
        }
    )

    for url in urls:
        try:
            fetch_url = _normalize_url_for_fetch(url)
            response = session.get(
                fetch_url,
                timeout=20,
                headers=_auth_headers_for_url(fetch_url, azure_devops_pat),
            )
            response.raise_for_status()
            text = _extract_response_text(response)
            if text:
                documents.append(
                    Document(
                        page_content=text,
                        metadata={"source": url, "fetch_url": fetch_url, "type": "url"},
                    )
                )
        except Exception as exc:  # pragma: no cover - defensive logging
            logger.warning("Failed to load url %s: %s", url, exc)
    return documents


def load_azure_devops_wiki_documents(
    organization: str,
    project: str,
    wiki_identifier: str,
    wiki_path: str,
    azure_devops_pat: str,
    api_version: str = "7.1",
) -> list[Document]:
    session = requests.Session()
    session.headers.update(
        {
            "Accept": "application/json",
            "User-Agent": "drs-chatbot-ingestion/1.0",
            **_basic_auth_headers(azure_devops_pat),
        }
    )
    endpoint = _azure_devops_wiki_endpoint(organization, project, wiki_identifier)
    response = session.get(
        endpoint,
        params={
            "path": _normalize_wiki_path(wiki_path),
            "recursionLevel": "full",
            "includeContent": "true",
            "api-version": api_version,
        },
        timeout=30,
    )
    _raise_for_azure_response(
        response,
        context=(
            f"Azure DevOps wiki read failed for project '{project}' and wiki '{wiki_identifier}'. "
            "Check that AZURE_DEVOPS_PAT is valid, has wiki read access, and can access this project."
        ),
    )
    payload = response.json()

    documents: list[Document] = []
    _collect_azure_wiki_pages(
        node=payload,
        documents=documents,
        session=session,
        endpoint=endpoint,
        api_version=api_version,
        project=project,
        wiki_identifier=wiki_identifier,
    )
    return documents


def load_all_azure_devops_project_wikis(
    organization: str,
    project: str,
    wiki_path: str,
    azure_devops_pat: str,
    api_version: str = "7.1",
) -> list[Document]:
    session = requests.Session()
    session.headers.update(
        {
            "Accept": "application/json",
            "User-Agent": "drs-chatbot-ingestion/1.0",
            **_basic_auth_headers(azure_devops_pat),
        }
    )
    response = session.get(
        _azure_devops_wikis_endpoint(organization, project),
        params={"api-version": api_version},
        timeout=30,
    )
    _raise_for_azure_response(
        response,
        context=(
            f"Azure DevOps wiki listing failed for project '{project}'. "
            "Check that AZURE_DEVOPS_PAT is valid, has the 'vso.wiki' scope, and can access this project. "
            "If you changed .env after starting the API, restart uvicorn so the new PAT is loaded."
        ),
    )
    payload = response.json()

    if isinstance(payload, list):
        wikis = payload
    else:
        wikis = payload.get("value") or []

    documents: list[Document] = []
    seen_identifiers: set[str] = set()
    for wiki in wikis:
        if not isinstance(wiki, dict):
            continue
        wiki_identifier = str(wiki.get("id") or wiki.get("name") or "").strip()
        if not wiki_identifier or wiki_identifier in seen_identifiers:
            continue
        seen_identifiers.add(wiki_identifier)
        try:
            documents.extend(
                load_azure_devops_wiki_documents(
                    organization=organization,
                    project=project,
                    wiki_identifier=wiki_identifier,
                    wiki_path=wiki_path,
                    azure_devops_pat=azure_devops_pat,
                    api_version=api_version,
                )
            )
        except Exception as exc:  # pragma: no cover - defensive logging
            logger.warning("Failed to load Azure DevOps wiki %s: %s", wiki_identifier, exc)
    return documents


def list_user_accessible_projects(
    organization: str,
    azure_devops_pat: str,
    api_version: str = "7.1",
) -> list[str]:
    org = quote(organization, safe="")
    response = requests.get(
        f"https://dev.azure.com/{org}/_apis/projects",
        params={"api-version": api_version, "stateFilter": "wellFormed"},
        headers={
            "Accept": "application/json",
            "User-Agent": "drs-chatbot-projects/1.0",
            **_basic_auth_headers(azure_devops_pat),
        },
        timeout=30,
    )
    _raise_for_azure_response(
        response,
        context=(
            f"Azure DevOps project listing failed for organization '{organization}'. "
            "Check that the supplied PAT is valid and has project-read access."
        ),
    )
    payload = response.json()
    items = payload if isinstance(payload, list) else (payload.get("value") or [])
    names: list[str] = []
    seen: set[str] = set()
    for entry in items:
        if not isinstance(entry, dict):
            continue
        name = str(entry.get("name") or "").strip()
        if name and name.lower() not in seen:
            seen.add(name.lower())
            names.append(name)
    return names


def load_sharepoint_documents(
    tenant_id: str,
    client_id: str,
    client_secret: str,
    site: str | None = None,
    folders: list[str] | None = None,
    graph_base_url: str = "https://graph.microsoft.com/v1.0",
    login_base_url: str = "https://login.microsoftonline.com",
    sink: IncrementalSink | None = None,
) -> list[Document]:
    """Ingest files from SharePoint via Microsoft Graph (app-only auth).

    Walks every site the app can read (or just ``site`` when provided), then every
    drive (document library) in each site, then recursively every folder, and loads
    each supported file. Each resulting Document carries ``folder``/``site``/``drive``
    metadata so answers stay traceable and so per-folder query scoping can be added
    later with the same pattern used for Azure DevOps ``project`` filtering.

    When ``sink`` is provided, each file is persisted the moment it loads (and files
    already in the store are skipped before download), so ingestion is resumable; the
    returned list is empty in that case.
    """
    token = _get_graph_token(tenant_id, client_id, client_secret, login_base_url)
    session = requests.Session()
    session.headers.update(
        {
            "Authorization": f"Bearer {token}",
            "Accept": "application/json",
            "User-Agent": "drs-chatbot-ingestion/1.0",
        }
    )
    base = graph_base_url.rstrip("/")

    folder_filter: list[str] | None = None
    if folders:
        cleaned = [f.strip() for f in folders if f and f.strip()]
        folder_filter = cleaned or None

    sites = _resolve_sharepoint_sites(session, base, site)

    documents: list[Document] = []
    seen_drives: set[str] = set()
    for site_obj in sites:
        if not isinstance(site_obj, dict):
            continue
        site_id = str(site_obj.get("id") or "").strip()
        if not site_id:
            continue
        site_name = str(site_obj.get("displayName") or site_obj.get("name") or site_id)
        try:
            drives = _graph_get_all(
                session,
                f"{base}/sites/{quote(site_id, safe='')}/drives",
                params={"$select": "id,name"},
                context=f"SharePoint drive listing failed for site '{site_name}'.",
            )
        except DocumentLoadError as exc:
            logger.warning("Skipping SharePoint site '%s': %s", site_name, exc)
            continue

        for drive in drives:
            if not isinstance(drive, dict):
                continue
            drive_id = str(drive.get("id") or "").strip()
            if not drive_id or drive_id in seen_drives:
                continue
            seen_drives.add(drive_id)
            drive_name = str(drive.get("name") or "")
            try:
                _walk_sharepoint_folder(
                    session=session,
                    base=base,
                    drive_id=drive_id,
                    item_id="root",
                    folder_path="",
                    site_name=site_name,
                    drive_name=drive_name,
                    folder_filter=folder_filter,
                    documents=documents,
                    sink=sink,
                )
            except DocumentLoadError as exc:
                logger.warning(
                    "Skipping SharePoint drive '%s' in site '%s': %s",
                    drive_name,
                    site_name,
                    exc,
                )
    return documents


def _get_graph_token(
    tenant_id: str,
    client_id: str,
    client_secret: str,
    login_base_url: str,
) -> str:
    url = f"{login_base_url.rstrip('/')}/{quote(tenant_id, safe='')}/oauth2/v2.0/token"
    response = requests.post(
        url,
        data={
            "client_id": client_id,
            "client_secret": client_secret,
            "scope": "https://graph.microsoft.com/.default",
            "grant_type": "client_credentials",
        },
        timeout=30,
    )
    if response.status_code >= 400:
        detail = ""
        try:
            detail = str(response.json().get("error_description") or "")
        except ValueError:
            detail = response.text[:300]
        raise DocumentLoadError(
            "SharePoint authentication failed "
            f"(HTTP {response.status_code}). {detail} "
            "Check SHAREPOINT_TENANT_ID, SHAREPOINT_CLIENT_ID, and SHAREPOINT_CLIENT_SECRET."
        )
    token = str(response.json().get("access_token") or "")
    if not token:
        raise DocumentLoadError("SharePoint authentication returned no access token.")
    return token


def _resolve_sharepoint_sites(
    session: requests.Session,
    base: str,
    site: str | None,
) -> list[dict]:
    if site and site.strip():
        value = site.strip()
        response = session.get(f"{base}/sites/{quote(value, safe=':/')}", timeout=30)
        if response.status_code < 400:
            return [response.json()]
        # Fall back to treating the value as a search term.
        results = _graph_get_all(
            session,
            f"{base}/sites",
            params={"search": value},
            context=f"SharePoint site search failed for '{value}'.",
        )
        if results:
            return results
        raise DocumentLoadError(
            f"SharePoint site '{value}' could not be resolved. "
            "Check SHAREPOINT_SITE and that the app has Sites.Read.All consent."
        )
    # No specific site: enumerate every site the app can read.
    return _graph_get_all(
        session,
        f"{base}/sites",
        params={"search": "*"},
        context="SharePoint site enumeration failed.",
    )


def _walk_sharepoint_folder(
    session: requests.Session,
    base: str,
    drive_id: str,
    item_id: str,
    folder_path: str,
    site_name: str,
    drive_name: str,
    folder_filter: list[str] | None,
    documents: list[Document],
    sink: IncrementalSink | None = None,
) -> None:
    children = _graph_get_all(
        session,
        f"{base}/drives/{quote(drive_id, safe='')}/items/{quote(item_id, safe='')}/children",
        params={"$select": "id,name,folder,file,webUrl,size", "$top": 200},
        context=f"SharePoint folder read failed for '{folder_path or '/'}'.",
    )
    for child in children:
        if not isinstance(child, dict):
            continue
        name = str(child.get("name") or "").strip()
        if not name:
            continue
        child_path = f"{folder_path}/{name}" if folder_path else f"/{name}"
        if isinstance(child.get("folder"), dict):
            _walk_sharepoint_folder(
                session=session,
                base=base,
                drive_id=drive_id,
                item_id=str(child.get("id") or ""),
                folder_path=child_path,
                site_name=site_name,
                drive_name=drive_name,
                folder_filter=folder_filter,
                documents=documents,
                sink=sink,
            )
        elif isinstance(child.get("file"), dict):
            document = _load_sharepoint_file(
                session=session,
                base=base,
                drive_id=drive_id,
                item=child,
                folder_path=folder_path,
                site_name=site_name,
                drive_name=drive_name,
                folder_filter=folder_filter,
                sink=sink,
            )
            if document is not None:
                if sink is not None:
                    sink.emit(document)
                else:
                    documents.append(document)


def _load_sharepoint_file(
    session: requests.Session,
    base: str,
    drive_id: str,
    item: dict,
    folder_path: str,
    site_name: str,
    drive_name: str,
    folder_filter: list[str] | None,
    sink: IncrementalSink | None = None,
) -> Document | None:
    name = str(item.get("name") or "").strip()
    if Path(name).suffix.lower() not in SUPPORTED_EXTENSIONS:
        return None
    if folder_filter is not None and not _folder_in_scope(folder_path, folder_filter):
        return None

    item_id = str(item.get("id") or "").strip()
    if not item_id:
        return None

    # Resume support: skip files already in the store before the expensive download.
    # This key must match the ``source`` metadata set on the Document below.
    normalized_folder = folder_path or "/"
    web_url = str(item.get("webUrl") or "")
    source_key = web_url or f"sharepoint:{site_name}:{normalized_folder}/{name}"
    if sink is not None and sink.is_ingested(source_key):
        logger.info("Skipping already-ingested SharePoint file: '%s'", name)
        return None

    logger.info(
        "Processing SharePoint file: site='%s' folder='%s' file='%s'",
        site_name,
        folder_path or "/",
        name,
    )
    response = session.get(
        f"{base}/drives/{quote(drive_id, safe='')}/items/{quote(item_id, safe='')}/content",
        timeout=60,
    )
    _raise_for_graph_response(
        response,
        context=f"SharePoint file download failed for '{name}'.",
    )
    text = _extract_file_text(name, response.content)
    if not text:
        return None

    return Document(
        page_content=text,
        metadata={
            "source": source_key,
            "folder": normalized_folder,
            "site": site_name,
            "drive": drive_name,
            "file_name": name,
            "type": "sharepoint",
        },
    )


def _folder_in_scope(folder_path: str, folder_filter: list[str]) -> bool:
    """True if folder_path matches a configured folder (as a contiguous segment)."""
    norm = "/" + "/".join(seg for seg in (folder_path or "").split("/") if seg)
    norm = norm.lower()
    for wanted in folder_filter:
        cleaned = "/".join(seg for seg in wanted.split("/") if seg).lower()
        if not cleaned:
            continue
        if norm == f"/{cleaned}" or f"/{cleaned}/" in f"{norm}/":
            return True
    return False


def _graph_get_all(
    session: requests.Session,
    url: str,
    params: dict | None = None,
    context: str = "Microsoft Graph request failed.",
) -> list[dict]:
    items: list[dict] = []
    next_url: str | None = url
    next_params = params
    while next_url:
        response = session.get(next_url, params=next_params, timeout=30)
        _raise_for_graph_response(response, context=context)
        payload = response.json()
        value = payload.get("value")
        if isinstance(value, list):
            items.extend(value)
        next_url = payload.get("@odata.nextLink")
        next_params = None  # nextLink already carries the query string.
    return items


def _raise_for_graph_response(response: requests.Response, context: str) -> None:
    if response.status_code < 400:
        return
    status = response.status_code
    if status in {401, 403}:
        raise DocumentLoadError(
            f"{context} Microsoft Graph returned HTTP {status}. "
            "Grant the app Sites.Read.All / Files.Read.All (application) permissions "
            "with admin consent."
        )
    raise DocumentLoadError(f"{context} Microsoft Graph returned HTTP {status}.")


def _extract_file_text(name: str, data: bytes) -> str:
    suffix = Path(name).suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf_bytes(data)
    if suffix in {".html", ".htm"}:
        return _extract_html_bytes(data)
    if suffix == ".docx":
        return _extract_docx_bytes(data)
    if suffix == ".xlsx":
        return _extract_xlsx_bytes(data)
    return _clean_text(data.decode("utf-8", errors="ignore"))


def _extract_docx_bytes(data: bytes) -> str:
    try:
        from docx import Document as DocxDocument
    except ImportError:  # pragma: no cover - optional dependency
        logger.warning("python-docx is not installed; skipping .docx file.")
        return ""
    try:
        document = DocxDocument(io.BytesIO(data))
    except Exception as exc:  # pragma: no cover - defensive logging
        logger.warning("Failed to parse .docx file: %s", exc)
        return ""

    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            line = " | ".join(cell for cell in cells if cell)
            if line:
                parts.append(line)
    return _clean_text("\n".join(parts))


def _extract_xlsx_bytes(data: bytes) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError:  # pragma: no cover - optional dependency
        logger.warning("openpyxl is not installed; skipping .xlsx file.")
        return ""
    try:
        workbook = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    except Exception as exc:  # pragma: no cover - defensive logging
        logger.warning("Failed to parse .xlsx file: %s", exc)
        return ""

    parts: list[str] = []
    total_chars = 0
    truncated = False
    try:
        for sheet in workbook.worksheets:
            if truncated:
                break
            rows_text: list[str] = []
            for row_index, row in enumerate(sheet.iter_rows(values_only=True)):
                if row_index >= MAX_XLSX_ROWS_PER_SHEET:
                    logger.warning(
                        "Sheet '%s' exceeded %d rows; truncating.",
                        sheet.title,
                        MAX_XLSX_ROWS_PER_SHEET,
                    )
                    truncated = True
                    break
                cells = [
                    str(value).strip()
                    for value in row
                    if value is not None and str(value).strip()
                ]
                if cells:
                    line = " | ".join(cells)
                    rows_text.append(line)
                    total_chars += len(line) + 1
                    if total_chars >= MAX_XLSX_CHARS:
                        logger.warning(
                            "Spreadsheet exceeded %d chars; truncating.",
                            MAX_XLSX_CHARS,
                        )
                        truncated = True
                        break
            if rows_text:
                parts.append(f"# Sheet: {sheet.title}\n" + "\n".join(rows_text))
    finally:
        workbook.close()
    return _clean_text("\n\n".join(parts))


def _extract_pdf_bytes(data: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(data))
        parts: list[str] = []
        for page in reader.pages:
            text = _clean_text(page.extract_text() or "")
            if text:
                parts.append(text)
        return "\n\n".join(parts)
    except Exception as exc:  # pragma: no cover - defensive logging
        logger.warning("Failed to parse SharePoint PDF: %s", exc)
        return ""


def _extract_html_bytes(data: bytes) -> str:
    html = data.decode("utf-8", errors="ignore")
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style", "noscript"]):
        tag.extract()
    return _clean_text(soup.get_text("\n"))


def split_documents(documents: list[Document], chunk_size: int, chunk_overlap: int) -> list[Document]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    chunks = splitter.split_documents(documents)
    for idx, doc in enumerate(chunks):
        doc.metadata["chunk_id"] = idx
    return chunks


def _load_single_file(path: Path) -> list[Document]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _load_pdf(path)
    if suffix in {".html", ".htm"}:
        return _load_html_file(path)
    if suffix == ".docx":
        text = _extract_docx_bytes(path.read_bytes())
        if not text:
            return []
        return [Document(page_content=text, metadata={"source": str(path), "type": "docx"})]
    if suffix == ".xlsx":
        text = _extract_xlsx_bytes(path.read_bytes())
        if not text:
            return []
        return [Document(page_content=text, metadata={"source": str(path), "type": "xlsx"})]
    return _load_text_file(path)


def _load_pdf(path: Path) -> list[Document]:
    documents: list[Document] = []
    reader = PdfReader(str(path))
    for page_idx, page in enumerate(reader.pages, start=1):
        text = _clean_text(page.extract_text() or "")
        if not text:
            continue
        documents.append(
            Document(
                page_content=text,
                metadata={"source": str(path), "page": page_idx, "type": "pdf"},
            )
        )
    return documents


def _load_html_file(path: Path) -> list[Document]:
    html = path.read_text(encoding="utf-8", errors="ignore")
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style", "noscript"]):
        tag.extract()
    text = _clean_text(soup.get_text("\n"))
    if not text:
        return []
    return [Document(page_content=text, metadata={"source": str(path), "type": "html"})]


def _load_text_file(path: Path) -> list[Document]:
    text = _clean_text(path.read_text(encoding="utf-8", errors="ignore"))
    if not text:
        return []
    return [Document(page_content=text, metadata={"source": str(path), "type": "text"})]


def _clean_text(text: str) -> str:
    lines = [line.strip() for line in text.splitlines()]
    return "\n".join(line for line in lines if line)


def _extract_response_text(response: requests.Response) -> str:
    content_type = (response.headers.get("Content-Type") or "").lower()
    if "text/html" in content_type or "<html" in response.text.lower():
        soup = BeautifulSoup(response.text, "html.parser")
        for tag in soup(["script", "style", "noscript"]):
            tag.extract()
        return _clean_text(soup.get_text("\n"))
    return _clean_text(response.text)


def _normalize_url_for_fetch(url: str) -> str:
    # Public Google Docs sharing links are best fetched via export endpoint.
    match = re.search(r"https?://docs\.google\.com/document/d/([a-zA-Z0-9_-]+)", url)
    if match:
        doc_id = match.group(1)
        return f"https://docs.google.com/document/d/{doc_id}/export?format=txt"
    return url


def _auth_headers_for_url(url: str, azure_devops_pat: str | None) -> dict[str, str]:
    if not azure_devops_pat:
        return {}

    host = urlparse(url).netloc.lower()
    if not _is_azure_devops_host(host):
        return {}

    return _basic_auth_headers(azure_devops_pat)


def _is_azure_devops_host(host: str) -> bool:
    return "dev.azure.com" in host or host.endswith(".visualstudio.com")


def _basic_auth_headers(personal_access_token: str) -> dict[str, str]:
    token = base64.b64encode(f":{personal_access_token}".encode("utf-8")).decode("ascii")
    return {"Authorization": f"Basic {token}"}


def _azure_devops_wiki_endpoint(organization: str, project: str, wiki_identifier: str) -> str:
    org = quote(organization, safe="")
    proj = quote(project, safe="")
    wiki = quote(wiki_identifier, safe="")
    return f"https://dev.azure.com/{org}/{proj}/_apis/wiki/wikis/{wiki}/pages"


def _azure_devops_wikis_endpoint(organization: str, project: str) -> str:
    org = quote(organization, safe="")
    proj = quote(project, safe="")
    return f"https://dev.azure.com/{org}/{proj}/_apis/wiki/wikis"


def _normalize_wiki_path(path: str | None) -> str:
    raw = (path or "/").strip()
    if not raw:
        return "/"
    return raw if raw.startswith("/") else f"/{raw}"


def _collect_azure_wiki_pages(
    node: dict,
    documents: list[Document],
    session: requests.Session,
    endpoint: str,
    api_version: str,
    project: str | None = None,
    wiki_identifier: str | None = None,
) -> None:
    content = _clean_text(str(node.get("content") or ""))
    page_path = str(node.get("path") or "")
    remote_url = str(node.get("remoteUrl") or "")
    page_id = node.get("id")

    if not content and page_path:
        content = _fetch_azure_wiki_page_content(
            session=session,
            endpoint=endpoint,
            page_path=page_path,
            api_version=api_version,
        )

    if content:
        documents.append(
            Document(
                page_content=content,
                metadata={
                    "source": remote_url or f"azure-wiki:{project or ''}:{page_path}",
                    "path": page_path,
                    "page_id": page_id,
                    "project": project,
                    "wiki": wiki_identifier,
                    "type": "azure_devops_wiki",
                },
            )
        )

    for child in node.get("subPages") or []:
        if isinstance(child, dict):
            _collect_azure_wiki_pages(
                node=child,
                documents=documents,
                session=session,
                endpoint=endpoint,
                api_version=api_version,
                project=project,
                wiki_identifier=wiki_identifier,
            )


def _fetch_azure_wiki_page_content(
    session: requests.Session,
    endpoint: str,
    page_path: str,
    api_version: str,
) -> str:
    response = session.get(
        endpoint,
        params={
            "path": _normalize_wiki_path(page_path),
            "includeContent": "true",
            "api-version": api_version,
        },
        timeout=30,
    )
    _raise_for_azure_response(
        response,
        context=f"Azure DevOps wiki page read failed for path '{page_path}'.",
    )
    payload = response.json()
    return _clean_text(str(payload.get("content") or ""))


def _raise_for_azure_response(response: requests.Response, context: str) -> None:
    try:
        response.raise_for_status()
    except requests.HTTPError as exc:
        status = response.status_code
        if status in {401, 403}:
            raise DocumentLoadError(context) from exc
        raise DocumentLoadError(f"{context} Azure returned HTTP {status}.") from exc
